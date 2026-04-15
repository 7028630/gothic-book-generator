import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import re
from PIL import Image, ImageDraw, ImageFont

# ============================================================
#  THE "NUCLEAR" CSS OVERRIDE - PRECISION REPAIR
# ============================================================

st.markdown("""
    <style>
    /* 1. Kill Top Bar and Decoration */
    header[data-testid="stHeader"], [data-testid="stDecoration"] {
        background: rgba(0,0,0,0) !important;
        background-color: transparent !important;
    }

    /* 2. Global Dark Theme */
    .stApp {
        background-color: #2b2b2b !important;
        color: #ffffff !important;
        font-family: 'Courier New', Courier, monospace !important;
    }

    /* 3. The Sidebar - Force White on Dark */
    [data-testid="stSidebar"] {
        background-color: #404040 !important;
    }
    [data-testid="stSidebar"] h3, 
    [data-testid="stSidebar"] label, 
    [data-testid="stSidebar"] p, 
    [data-testid="stSidebar"] span {
        color: #ffffff !important;
    }

    /* 4. EXPANDER FIX: Anti-Whiteout Logic */
    div[data-testid="stExpander"] {
        background-color: transparent !important;
        border: 1px solid #696969 !important;
    }
    div[data-testid="stExpander"] details {
        background-color: transparent !important;
    }
    div[data-testid="stExpander"] details summary {
        background-color: transparent !important;
        color: #ffffff !important;
    }
    div[data-testid="stExpander"] details[open] > summary {
        background-color: transparent !important;
        color: #ffffff !important;
        border-bottom: 1px solid #696969;
    }
    div[data-testid="stExpander"] [data-testid="stVerticalBlock"] {
        background-color: transparent !important;
    }

    /* 5. CODE BLOCKS / COPY BUTTONS: Gray Scale Fix */
    code {
        color: #ffffff !important;
        background-color: #404040 !important; /* Dark Gray */
        border: 1px solid #696969 !important;
    }
    div[data-testid="stCodeBlock"] {
        background-color: #404040 !important;
    }
    /* Target the copy button specifically */
    div[data-testid="stCodeBlock"] button {
        background-color: #000000 !important;
        color: white !important;
    }

    /* 6. Uploader Section */
    [data-testid="stFileUploader"] section {
        background-color: #333333 !important;
        border: 1px dashed #000000 !important;
        color: #000000 !important;
    }

    /* 7. Main Buttons */
    button[kind="secondary"], button[kind="primary"] {
        background-color: #696969 !important;
        border: 2px solid #ffffff !important;
        border-radius: 0px !important;
    }
    button p {
        color: #000000 !important;
        font-weight: 900 !important;
    }
    button:hover {
        background-color: #ffffff !important;
    }
    </style>
    """, unsafe_allow_html=True)

# ============================================================
#  ASSET GENERATORS
# ============================================================

def get_gothic_asset(text, color_rgb, font_size=80):
    try:
        font = ImageFont.truetype("Friedolin.ttf", font_size)
    except:
        font = ImageFont.load_default()
    
    left, top, right, bottom = font.getbbox(text)
    img = Image.new('RGBA', (right-left + 60, bottom-top + 40), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    draw.text((30, 10), text, font=font, fill=color_rgb)
    
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf

def add_floating_element(doc, img_buf, width_cm, x_cm, y_cm):
    run = doc.add_paragraph().add_run()
    shape = run.add_picture(img_buf, width=Cm(width_cm))
    inline = shape._inline
    anchor = OxmlElement('wp:anchor')
    anchor.set('distT', '0'); anchor.set('distB', '0'); anchor.set('distL', '0'); anchor.set('distR', '0')
    anchor.set('simplePos', '0'); anchor.set('relativeHeight', '251658240')
    anchor.set('behindDoc', '0'); anchor.set('locked', '0')
    anchor.set('layoutInCell', '1'); anchor.set('allowOverlap', '1')
    
    posH = OxmlElement('wp:positionH'); posH.set('relativeFrom', 'page')
    posH.append(OxmlElement('wp:posOffset'))
    posH.find(qn('wp:posOffset')).text = str(int(x_cm * 360000))
    
    posV = OxmlElement('wp:positionV'); posV.set('relativeFrom', 'page')
    posV.append(OxmlElement('wp:posOffset'))
    posV.find(qn('wp:posOffset')).text = str(int(y_cm * 360000))
    
    anchor.append(posH); anchor.append(posV)
    for child in inline:
        if child.tag != qn('wp:docPr') and child.tag != qn('wp:cNvGraphicFramePr'):
            anchor.append(child)
    inline.getparent().replace(inline, anchor)

# ============================================================
#  MAIN APP
# ============================================================

def main():
    try:
        title_png = get_gothic_asset("Gothic Book Generator", (255, 255, 255), 100)
        st.image(title_png)
    except:
        st.title("Gothic Book Generator")

    if 'img_lib' not in st.session_state: 
        st.session_state.img_lib = {}

    with st.sidebar:
        st.markdown("### 🎨 Colors")
        t_color = st.color_picker("Title Color", "#8B0000")
        s_color = st.color_picker("Subtitle Color", "#FFFFFF")
        rgb_title = tuple(int(t_color.lstrip('#')[i:i+2], 16) for i in (0, 2, 4))
        rgb_sub = tuple(int(s_color.lstrip('#')[i:i+2], 16) for i in (0, 2, 4))

        st.divider()
        st.markdown("### 🖋️ Typography")
        main_size = st.slider("Main Text Size", 8, 30, 12)
        note_size = main_size - 1
        note_font = "Courier New"

        st.divider()
        st.markdown("### 🖼️ Illustrations")
        uploads = st.file_uploader("Upload Assets", accept_multiple_files=True)
        if uploads:
            for up in uploads:
                st.session_state.img_lib[up.name] = up
                st.write(f"Ref: {up.name}")
                st.code(f"[IMG: {up.name}]")

    with st.expander("📖 HOW TO COMPOSE YOUR TEXT"):
        st.markdown(f"""
        **Commands for your .txt file:**
        * `[TITLE: Text]` → Gothic Title (using Title Color).
        * `[SUB: Text]` → Gothic Subtitle (using Subtitle Color).
        * `[IMG: filename.png]` → Inserts an illustration.
        * `[NOTE_START]` → Inserts `Separator.png` and starts **{note_size}pt** text.
        * `[NOTE_END]` → Inserts `Separator.png` and returns to **{main_size}pt** text.
        * **Standard Text** → Formatted at **{main_size}pt**.
        """)

    st.write("🏛️ **UPLOAD MAIN TEXT (.TXT)**")
    notepads = st.file_uploader("Main File", accept_multiple_files=True, label_visibility="collapsed")

    if notepads and st.button("🚀 Build A4 Horizontal Book"):
        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = Cm(29.7), Cm(21.0)
        section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Cm(1.5)

        pg_num = 1
        for note in notepads:
            lines = note.read().decode("utf-8").split('\n')
            table = doc.add_table(rows=2, cols=2)
            table.autofit = False
            cell_l = table.rows[0].cells[0]
            current_y = 2.0
            in_commentary = False
            
            for line in lines:
                line = line.strip()
                if not line: continue
                
                if line == "[NOTE_START]":
                    in_commentary = True
                    if "Separator.png" in st.session_state.img_lib:
                        add_floating_element(doc, st.session_state.img_lib["Separator.png"], 12, 1.5, current_y)
                        current_y += 1.5
                    continue
                
                if line == "[NOTE_END]":
                    in_commentary = False
                    if "Separator.png" in st.session_state.img_lib:
                        add_floating_element(doc, st.session_state.img_lib["Separator.png"], 12, 1.5, current_y)
                        current_y += 1.5
                    continue

                if line.startswith("[TITLE:"):
                    txt = re.search(r"\[TITLE: (.*?)\]", line).group(1)
                    add_floating_element(doc, get_gothic_asset(txt, rgb_title, 80), 9, 2, current_y)
                    current_y += 3.0
                    for _ in range(4): cell_l.add_paragraph()

                elif line.startswith("[SUB:"):
                    txt = re.search(r"\[SUB: (.*?)\]", line).group(1)
                    add_floating_element(doc, get_gothic_asset(txt, rgb_sub, 50), 6, 2, current_y)
                    current_y += 2.0
                    for _ in range(2): cell_l.add_paragraph()
                
                elif line.startswith("[IMG:"):
                    name = re.search(r"\[IMG: (.*?)\]", line).group(1)
                    if name in st.session_state.img_lib:
                        add_floating_element(doc, st.session_state.img_lib[name], 6, 3, current_y)
                        current_y += 6.5
                else:
                    p = cell_l.add_paragraph(line)
                    run = p.runs[0] if p.runs else p.add_run(line)
                    run.font.name = note_font
                    run.font.size = Pt(note_size) if in_commentary else Pt(main_size)
                    p.paragraph_format.first_line_indent = 0
            
            table.rows[1].cells[0].add_paragraph(str(pg_num)).alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.rows[1].cells[1].add_paragraph(str(pg_num + 1)).alignment = WD_ALIGN_PARAGRAPH.CENTER
            pg_num += 2
            doc.add_page_break()

        out = io.BytesIO()
        doc.save(out)
        st.download_button("📥 Download Book", out.getvalue(), "gothic_spreads.docx")

if __name__ == "__main__":
    main()
